import json
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Style parameters
FONT_TITLE = 'Georgia'
FONT_BODY = 'Arial'
COLOR_VIOLET = RGBColor(76, 43, 115)     # Deep luxury violet
COLOR_CHARCOAL = RGBColor(44, 44, 44)    # Dark charcoal body
COLOR_MUTED = RGBColor(110, 110, 110)    # Muted gray

def clean_html_text(text):
    if not text:
        return ""
    text = text.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
    text = text.replace("<em>", "").replace("</em>", "")
    text = text.replace("<strong>", "").replace("</strong>", "")
    text = text.replace("&copy;", "©")
    return " ".join(text.split())

def add_docx_page_content(doc, title, subtitle, blocks):
    # Add page title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(title.upper())
    run_title.font.name = FONT_TITLE
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_VIOLET
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(4)

    if subtitle:
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = FONT_TITLE
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.color.rgb = COLOR_MUTED
        p_sub.paragraph_format.space_after = Pt(16)

    for block in blocks:
        heading = block['heading']
        items = block['items']
        
        if heading:
            p_head = doc.add_paragraph()
            run_head = p_head.add_run(heading)
            run_head.font.name = FONT_TITLE
            run_head.font.size = Pt(13)
            run_head.font.bold = True
            run_head.font.color.rgb = COLOR_VIOLET
            p_head.paragraph_format.space_before = Pt(14)
            p_head.paragraph_format.space_after = Pt(8)
            p_head.paragraph_format.keep_with_next = True
            
        for item_type, data in items:
            if item_type == 'card':
                title_val, desc_val = data
                p = doc.add_paragraph()
                
                # Title part (bold)
                if title_val:
                    run_title = p.add_run(title_val)
                    run_title.font.name = FONT_BODY
                    run_title.font.size = Pt(10)
                    run_title.font.bold = True
                    run_title.font.color.rgb = COLOR_CHARCOAL
                    
                # If both title and desc are present, add a line break
                if title_val and desc_val:
                    p.add_run("\n")
                    
                # Description part
                if desc_val:
                    run_desc = p.add_run(desc_val)
                    run_desc.font.name = FONT_BODY
                    run_desc.font.size = Pt(10)
                    run_desc.font.color.rgb = COLOR_CHARCOAL
                    
                p.paragraph_format.space_after = Pt(8)
                
            elif item_type == 'subheading':
                p = doc.add_paragraph()
                run = p.add_run(data.upper())
                run.font.name = FONT_TITLE
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = COLOR_VIOLET
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                
            elif item_type == 'paragraph':
                p = doc.add_paragraph()
                run = p.add_run(data)
                run.font.name = FONT_BODY
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_CHARCOAL
                p.paragraph_format.space_after = Pt(6)
                
            elif item_type == 'quote':
                p = doc.add_paragraph()
                run = p.add_run(data)
                run.font.name = FONT_BODY
                run.font.size = Pt(9.5)
                run.font.italic = True
                run.font.color.rgb = COLOR_MUTED
                p.paragraph_format.space_after = Pt(6)

def parse_html_page_to_blocks(html_path, de_dict):
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    # Decompose Header, Footer, and irrelevant elements
    for el in soup.find_all(['header', 'footer', 'script', 'style', 'noscript']):
        el.decompose()
    for el in soup.find_all(class_=['mobile-menu', 'lang-float', 'cookie-consent', 'back-to-top', 'lang-float-dropdown', 'anchor-nav']):
        el.decompose()

    main_el = soup.find('main') or soup.find('body')
    if not main_el:
        return []

    # Find top-level sections
    sections = main_el.find_all('section', recursive=False)
    if not sections:
        # Fallback to general sections
        sections = main_el.find_all('section')
    if not sections:
        sections = [main_el]

    blocks = []
    seen_elements = set()

    for sec in sections:
        if not sec:
            continue
        
        # 1. Gather Section Heading
        heading_text = ""
        headings = sec.find_all(['h1', 'h2', 'h3'])
        for h in headings:
            h_class = h.get('class') or []
            if 'section-tag' in h_class or 'service-group-label' in h_class:
                continue
            
            h_key = h.get('data-i18n') or h.get('data-i18n-ph')
            h_val = de_dict.get(h_key) if h_key else None
            if not h_val:
                h_val = h.text.strip()
            
            h_clean = clean_html_text(h_val)
            # Filter out generic button texts or small headers
            if h_clean and not any(x in h_clean.lower() for x in ['anrufen', 'kontakt', 'beratung', 'bewerben']):
                heading_text = h_clean
                seen_elements.add(id(h))
                break

        # 2. Extract content items
        sec_items = []
        
        def get_el_text(el):
            key = el.get('data-i18n') or el.get('data-i18n-ph')
            val = de_dict.get(key) if key else None
            if not val:
                val = el.text.strip()
            return clean_html_text(val)

        # Find structured containers inside the section
        containers = sec.find_all(class_=lambda c: c and any(x in c for x in ['card', 'process-step', 'fort-item', 'kosten-item', 'stat-item', 'review-card', 'kosten-card']))
        lis = sec.find_all('li')
        
        # Combine in natural preorder sequence
        all_containers = []
        for descendant in sec.descendants:
            if descendant in containers or descendant in lis:
                if descendant not in all_containers:
                    all_containers.append(descendant)
        
        processed_elements = set()

        if all_containers:
            for cont in all_containers:
                if id(cont) in processed_elements:
                    continue
                processed_elements.add(id(cont))
                
                # Filter buttons
                cont_class = cont.get('class') or []
                if 'btn' in cont_class or any('btn' in x for x in cont_class):
                    continue

                # Find title & description
                title_el = cont.find(['h3', 'h4', 'strong'])
                title_val = ""
                if title_el:
                    title_val = get_el_text(title_el)
                    processed_elements.add(id(title_el))

                desc_el = cont.find(['p', 'span'])
                if desc_el == title_el:
                    desc_el = None
                desc_val = ""
                if desc_el:
                    desc_val = get_el_text(desc_el)
                    processed_elements.add(id(desc_el))

                # If it's a simple list item / card with no structured children, grab direct texts
                if not title_val and not desc_val:
                    # Traverse texts inside container
                    text_parts = []
                    for desc in cont.descendants:
                        if not desc.name and desc.string and desc.string.strip():
                            p_el = desc.parent
                            if p_el.name in ['p', 'span', 'strong', 'em', 'h3', 'h4', 'div', 'li']:
                                if not p_el.get('data-i18n') and not p_el.get('data-i18n-ph'):
                                    txt = clean_html_text(desc.string.strip())
                                    if len(txt) > 2 and txt not in text_parts:
                                        text_parts.append(txt)
                                p_key = p_el.get('data-i18n') or p_el.get('data-i18n-ph')
                                if p_key and de_dict.get(p_key):
                                    txt = clean_html_text(de_dict[p_key])
                                    if len(txt) > 2 and txt not in text_parts:
                                        text_parts.append(txt)
                    
                    if len(text_parts) >= 2:
                        title_val = text_parts[0]
                        desc_val = " ".join(text_parts[1:])
                    elif len(text_parts) == 1:
                        title_val = text_parts[0]

                # Filter buttons or boilerplate
                if title_val and any(x in title_val.lower() for x in ['anrufen', 'kontakt', 'beratung', 'bewerben', 'mehr über', 'zur hauptseite']):
                    continue
                if desc_val and any(x in desc_val.lower() for x in ['anrufen', 'kontakt', 'beratung', 'bewerben', 'mehr über', 'zur hauptseite']):
                    continue

                if title_val or desc_val:
                    sec_items.append(('card', (title_val, desc_val)))

        # Also get independent text nodes that were not processed in containers
        for el in sec.find_all(['p', 'div', 'cite']):
            # Verify it's not nested inside a processed container
            is_processed = id(el) in processed_elements
            if not is_processed:
                for parent in el.parents:
                    if id(parent) in processed_elements:
                        is_processed = True
                        break
            
            if is_processed:
                continue
                
            el_class = el.get('class') or []
            if 'btn' in el_class or any('btn' in x for x in el_class) or 'section-tag' in el_class:
                continue
                
            if el.name == 'p' or el.name == 'div' or el.name == 'cite':
                if 'service-group-label' in el_class or 'group-label' in el_class:
                    lbl = get_el_text(el)
                    if lbl:
                        sec_items.append(('subheading', lbl))
                    continue
                
                if 'quote-block' in el_class or el.name == 'cite' or any('quote' in x for x in el_class):
                    txt = get_el_text(el)
                    if txt:
                        sec_items.append(('quote', txt))
                    continue
                
                # Check normal paragraph
                if el.name == 'p' or 'accent' in el_class or 'lead' in el_class:
                    txt = get_el_text(el)
                    if txt and not any(x in txt.lower() for x in ['anrufen', 'kontakt', 'beratung', 'bewerben', 'mehr über', 'zur hauptseite']):
                        sec_items.append(('paragraph', txt))

        # Check for heading duplicate in cards
        clean_items = []
        for item_type, data in sec_items:
            if item_type == 'card':
                title_val, desc_val = data
                if title_val == heading_text:
                    if desc_val:
                        clean_items.append(('paragraph', desc_val))
                    continue
            clean_items.append((item_type, data))

        if clean_items or heading_text:
            blocks.append({
                'heading': heading_text,
                'items': clean_items
            })

    return blocks

def generate_docs():
    corp_site_dir = "/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-farmos/site"
    corp_de_json = os.path.join(corp_site_dir, 'i18n/de.json')
    output_corp_docx = "/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-farmos/site/сайт-текст/KSK-Farmos-Website-Text-DE.docx"

    print("Generating Corporate Website Docx...")
    if os.path.exists(corp_de_json):
        with open(corp_de_json, 'r', encoding='utf-8') as f:
            corp_de_dict = json.load(f)

        doc_corp = Document()
        for section in doc_corp.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Title Page
        p_spacer = doc_corp.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(80)
        
        p_title = doc_corp.add_paragraph()
        run_title = p_title.add_run("KSK Farmos GmbH & Co. KG")
        run_title.font.name = FONT_TITLE
        run_title.font.size = Pt(26)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_VIOLET
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_sub = doc_corp.add_paragraph()
        run_sub = p_sub.add_run("Inhalt der Website (Deutsch / Original)")
        run_sub.font.name = FONT_TITLE
        run_sub.font.size = Pt(13)
        run_sub.font.italic = True
        run_sub.font.color.rgb = COLOR_CHARCOAL
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(200)

        p_meta = doc_corp.add_paragraph()
        run_meta = p_meta.add_run("Dokumententyp: Website-Textgliederung (ohne Menü und Fußzeile)\nGeneriert am: 18. Mai 2026\nIntensivpflegedienst Nordhessen")
        run_meta.font.name = FONT_BODY
        run_meta.font.size = Pt(9.5)
        run_meta.font.color.rgb = COLOR_MUTED
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_corp.add_page_break()

        # Pages in logical order
        corp_pages = [
            ('index.html', 'Startseite', 'Hauptseite des Unternehmens'),
            ('leistungen.html', 'Leistungen', 'Medizinische Leistungen & Pflegekonzepte'),
            ('ueber-uns.html', 'Über uns', 'Unternehmensphilosophie & Team'),
            ('karriere.html', 'Karriere', 'Stellenangebote & Arbeitgeberleistungen'),
            ('schnellbewerbung.html', 'Schnellbewerbung', 'Interaktiver Bewerbungsprozess'),
            ('faq.html', 'FAQ', 'Häufig gestellte Fragen'),
            ('kontakt.html', 'Kontakt', 'Standorte & Kontaktaufnahme'),
            ('beratung.html', 'Beratung', 'Kostenloses Erstgespräch anfordern')
        ]

        for filename, title, subtitle in corp_pages:
            filepath = os.path.join(corp_site_dir, filename)
            if os.path.exists(filepath):
                print(f"Parsing {filename}...")
                blocks = parse_html_page_to_blocks(filepath, corp_de_dict)
                if blocks:
                    add_docx_page_content(doc_corp, title, subtitle, blocks)
                    doc_corp.add_page_break()

        doc_corp.save(output_corp_docx)
        print(f"Corporate document created successfully: {output_corp_docx}")
    else:
        print(f"Warning: Corporate de.json not found at {corp_de_json}")

    # ─────────────────────────────────────────────────────────
    # DOCUMENT 2: KSK PATIENTS LANDING PAGE (GERMAN)
    # ─────────────────────────────────────────────────────────
    pat_site_dir = "/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-patients/landing-intensivpflege"
    pat_de_json = os.path.join(pat_site_dir, 'i18n/de.json')
    output_pat_docx = "/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-farmos/site/сайт-текст/KSK-Patients-Landing-Text-DE.docx"

    print("Generating Patients Landing Page Docx...")
    if os.path.exists(pat_de_json):
        with open(pat_de_json, 'r', encoding='utf-8') as f:
            pat_de_dict = json.load(f)

        doc_pat = Document()
        for section in doc_pat.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Title Page
        p_spacer = doc_pat.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(80)
        
        p_title = doc_pat.add_paragraph()
        run_title = p_title.add_run("KSK Farmos — Intensivpflege")
        run_title.font.name = FONT_TITLE
        run_title.font.size = Pt(26)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_VIOLET
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_sub = doc_pat.add_paragraph()
        run_sub = p_sub.add_run("Inhalt der Patienten-Landingpage (Deutsch / Original)")
        run_sub.font.name = FONT_TITLE
        run_sub.font.size = Pt(13)
        run_sub.font.italic = True
        run_sub.font.color.rgb = COLOR_CHARCOAL
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(200)

        p_meta = doc_pat.add_paragraph()
        run_meta = p_meta.add_run("Dokumententyp: Landingpage-Textgliederung (ohne Menü und Fußzeile)\nGeneriert am: 18. Mai 2026\nPatienten-Werbelandingpage Nordhessen")
        run_meta.font.name = FONT_BODY
        run_meta.font.size = Pt(9.5)
        run_meta.font.color.rgb = COLOR_MUTED
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_pat.add_page_break()

        pat_filepath = os.path.join(pat_site_dir, 'index.html')
        if os.path.exists(pat_filepath):
            print(f"Parsing landing page index.html...")
            blocks = parse_html_page_to_blocks(pat_filepath, pat_de_dict)
            if blocks:
                add_docx_page_content(doc_pat, 'Patienten-Landingpage', 'Ambulante Intensivpflege Nordhessen', blocks)

        doc_pat.save(output_pat_docx)
        print(f"Patients Landing document created successfully: {output_pat_docx}")
    else:
        print(f"Warning: Patients de.json not found at {pat_de_json}")

if __name__ == '__main__':
    generate_docs()
