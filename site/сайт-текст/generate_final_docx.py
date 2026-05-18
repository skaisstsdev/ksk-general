import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Styling constants
FONT_TITLE = 'Georgia'
FONT_BODY = 'Arial'
COLOR_VIOLET = RGBColor(76, 43, 115)     # Deep luxury violet
COLOR_CHARCOAL = RGBColor(44, 44, 44)    # Dark charcoal body
COLOR_MUTED = RGBColor(110, 110, 110)    # Muted gray

def build_docx(filepath, doc_title, doc_subtitle, metadata_text, pages):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. Cover Page
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(80)
    
    p_main_title = doc.add_paragraph()
    run_main_title = p_main_title.add_run(doc_title)
    run_main_title.font.name = FONT_TITLE
    run_main_title.font.size = Pt(26)
    run_main_title.font.bold = True
    run_main_title.font.color.rgb = COLOR_VIOLET
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(doc_subtitle)
    run_sub.font.name = FONT_TITLE
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_CHARCOAL
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(200)

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(metadata_text)
    run_meta.font.name = FONT_BODY
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = COLOR_MUTED
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2. Content Pages
    for page_idx, page in enumerate(pages):
        page_title = page['title']
        page_desc = page.get('desc', '')
        blocks = page['blocks']

        # Add page title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(page_title.upper())
        run_title.font.name = FONT_TITLE
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_VIOLET
        p_title.paragraph_format.space_before = Pt(24)
        p_title.paragraph_format.space_after = Pt(4)

        if page_desc:
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run(page_desc)
            run_sub.font.name = FONT_TITLE
            run_sub.font.size = Pt(11)
            run_sub.font.italic = True
            run_sub.font.color.rgb = COLOR_MUTED
            p_sub.paragraph_format.space_after = Pt(16)

        # Add page blocks
        for block in blocks:
            heading = block.get('heading', '')
            items = block.get('items', [])

            if heading:
                p_head = doc.add_paragraph()
                run_head = p_head.add_run(heading)
                run_head.font.name = FONT_TITLE
                run_head.font.size = Pt(13)
                run_head.font.bold = True
                run_head.font.color.rgb = COLOR_VIOLET
                p_head.paragraph_format.space_before = Pt(14)
                p_head.paragraph_format.space_after = Pt(8)
                p_head.paragraph_format.keep_with_next = True

            for item in items:
                item_type = item[0]
                
                if item_type == 'paragraph':
                    text = item[1]
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLOR_CHARCOAL
                    p.paragraph_format.space_after = Pt(6)
                    
                elif item_type == 'card':
                    title, desc = item[1], item[2]
                    p = doc.add_paragraph()
                    
                    if title:
                        run_t = p.add_run(title)
                        run_t.font.name = FONT_BODY
                        run_t.font.size = Pt(10)
                        run_t.font.bold = True
                        run_t.font.color.rgb = COLOR_CHARCOAL
                        
                    if title and desc:
                        p.add_run("\n")
                        
                    if desc:
                        run_d = p.add_run(desc)
                        run_d.font.name = FONT_BODY
                        run_d.font.size = Pt(10)
                        run_d.font.color.rgb = COLOR_CHARCOAL
                        
                    p.paragraph_format.space_after = Pt(8)

                elif item_type == 'subheading':
                    text = item[1]
                    p = doc.add_paragraph()
                    run = p.add_run(text.upper())
                    run.font.name = FONT_TITLE
                    run.font.size = Pt(10.5)
                    run.font.bold = True
                    run.font.color.rgb = COLOR_VIOLET
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.keep_with_next = True
                    
                elif item_type == 'bullet':
                    text = item[1]
                    p = doc.add_paragraph()
                    run = p.add_run(f"•  {text}")
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLOR_CHARCOAL
                    p.paragraph_format.space_after = Pt(4)

                elif item_type == 'quote':
                    text = item[1]
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9.5)
                    run.font.italic = True
                    run.font.color.rgb = COLOR_MUTED
                    p.paragraph_format.space_after = Pt(6)

        # Add page break except for last page
        if page_idx < len(pages) - 1:
            doc.add_page_break()

    doc.save(filepath)
    print(f"Successfully generated: {filepath}")

# ─────────────────────────────────────────────────────────
# DATA 1: KSK FARMOS CORPORATE WEBSITE
# ─────────────────────────────────────────────────────────
corp_pages = [
    {
        'title': 'Startseite',
        'desc': 'Hauptseite des Unternehmens',
        'blocks': [
            {
                'heading': 'Professionelle Pflege, die ankommt.',
                'items': [
                    ('paragraph', 'Ambulante Intensivpflege bei Ihnen zu Hause oder in unserer Einrichtung in Kassel — 24 Stunden am Tag, 7 Tage die Woche.'),
                    ('paragraph', 'Seit 2013 versorgen wir Menschen mit schweren Erkrankungen in ganz Hessen. Mit durchschnittlich 5,5 Mitarbeitern pro Patient und einem Team aus examinierten Fachkräften garantieren wir eine 24/7-Versorgung auf höchstem Niveau — im eigenen Zuhause oder in unserer Einrichtung in Kassel.')
                ]
            },
            {
                'heading': 'Wie wir helfen',
                'items': [
                    ('card', 'Häusliche Intensivpflege', 'Professionelle Versorgung in Ihren eigenen vier Wänden — rund um die Uhr.'),
                    ('card', 'Aufenthaltskonzept Kassel', 'Eigenes Zimmer, Garten, Bibliothek und 24h-Betreuung an der Sommerbergstraße 14.'),
                    ('card', 'Fort- und Weiterbildungen', 'Regelmäßige Schulungen durch Fachärzte in 8 medizinischen Fachbereichen.')
                ]
            },
            {
                'heading': 'Wen wir versorgen',
                'items': [
                    ('card', 'Beatmungspatienten', 'Invasive und nicht-invasive Beatmung'),
                    ('card', 'Wachkoma (Apallisches Syndrom)', 'Langzeitversorgung und Stimulation'),
                    ('card', 'ALS / Neurologische Erkrankungen', 'Progressive Begleitung'),
                    ('card', 'Querschnittslähmung', 'Ganzheitliche Pflege'),
                    ('card', 'Schädel-Hirn-Trauma', 'Rehabilitation und Stabilisierung'),
                    ('quote', '"Das Wohl des Patienten ist oberstes Gesetz." — Unser Leitbild')
                ]
            },
            {
                'heading': 'In 3 Schritten zur Versorgung',
                'items': [
                    ('card', 'Erstes Gespräch', 'Kostenlos und unverbindlich — wir lernen Ihre Situation kennen und beraten Sie persönlich.'),
                    ('card', 'Individuelle Planung', 'Wir erstellen einen Pflegeplan, stimmen uns mit Ärzten ab und klären die Kostenübernahme.'),
                    ('card', 'Versorgungsbeginn', 'Ihr persönliches Pflegeteam beginnt die Versorgung — professionell und zuverlässig.')
                ]
            },
            {
                'heading': 'Wer wir sind',
                'items': [
                    ('paragraph', 'Der ambulante Intensivpflegedienst KSK Farmos GmbH & Co. KG wurde 2013 von Viktor Beresnev in Volkmarsen gegründet. Seitdem versorgen wir Patienten hessenweit — mit Herz, Fachwissen und modernster Ausstattung.'),
                    ('paragraph', '5,5 Mitarbeiter pro Patient — weit über dem Branchendurchschnitt.')
                ]
            },
            {
                'heading': 'Das sagen unsere Patienten',
                'items': [
                    ('quote', '"Sehr kompetentes und einfühlsames Team. Die 24h-Betreuung in Kassel ist eine enorme Entlastung für unsere Familie. Wir fühlen uns perfekt aufgehoben."\n— Familie M., Angehörige, Kassel'),
                    ('quote', '"Dank der häuslichen Intensivpflege von KSK Farmos GmbH & Co. KG kann mein Vater in seiner gewohnten Umgebung bleiben. Zuverlässig und stets freundlich."\n— Thomas W., Angehöriger, Volkmarsen'),
                    ('quote', '"Vom ersten Gespräch bis zur Umsetzung der Pflege lief alles reibungslos. Die Pflegekräfte sind hochqualifiziert und arbeiten sehr professionell."\n— Sabine K., Patientin, Korbach')
                ]
            }
        ]
    },
    {
        'title': 'Leistungen',
        'desc': 'Medizinische Leistungen & Pflegekonzepte',
        'blocks': [
            {
                'heading': 'Unsere medizinischen Leistungen',
                'items': [
                    ('paragraph', 'Intensivpflege zu Hause, einzigartige Wohnkonzepte und ständige Weiterbildung — alles aus einer Hand.')
                ]
            },
            {
                'heading': 'Häusliche Intensivpflege',
                'items': [
                    ('paragraph', 'Wir versorgen Menschen mit schweren Erkrankungen in ihrem gewohnten Umfeld — 24 Stunden am Tag, 7 Tage die Woche, 365 Tage im Jahr.'),
                    ('subheading', 'Grundversorgung'),
                    ('bullet', 'Grund- und Behandlungspflege'),
                    ('bullet', 'Medikamentenmanagement'),
                    ('bullet', 'Vitalzeichenüberwachung und Monitoring'),
                    ('subheading', 'Spezielle Therapien'),
                    ('bullet', 'Invasive und nicht-invasive Beatmung'),
                    ('bullet', 'Trachealkanülenmanagement'),
                    ('bullet', 'Wundversorgung und Ernährungstherapie'),
                    ('subheading', 'Organisation'),
                    ('bullet', 'Pflegeplanung und Dokumentation'),
                    ('bullet', 'Koordination mit Ärzten und Therapeuten'),
                    ('quote', 'Die Beatmungspflege ist unsere Kernkompetenz. Wir betreuen Patienten mit invasiver und nicht-invasiver Beatmung — mit erfahrenen Fachkräften, modernster Technik und in enger Zusammenarbeit mit Pneumologen und Anästhesisten.')
                ]
            },
            {
                'heading': 'Wen wir versorgen',
                'items': [
                    ('card', 'Beatmungspatienten', 'Spezialisierte Intensivpflege für Menschen, die auf invasive oder nicht-invasive Beatmung angewiesen sind. Wir sichern die lückenlose Überwachung und professionelles Gerätemanagement.'),
                    ('card', 'Wachkoma (Apallisches Syndrom)', 'Ganzheitliche Langzeitversorgung für Patienten mit schweren Bewusstseinsstörungen. Unser Fokus liegt auf basaler Stimulation und der Erhaltung von Körperfunktionen.'),
                    ('card', 'ALS / Neurologische Erkrankungen', 'Individuelle Begleitung bei fortschreitenden neurologischen Veränderungen — empathisch, fachlich fundiert und auf die Erhaltung der Lebensqualität ausgerichtet.'),
                    ('card', 'Querschnittslähmung', 'Professionelle Unterstützung und Pflege für Menschen mit Rückenmarksverletzungen. Wir fördern die Selbstständigkeit und beugen gezielt Komplikationen vor.'),
                    ('card', 'Schädel-Hirn-Trauma', 'Langfristige Versorgung nach schweren Hirnverletzungen. Zielgerichtet auf die Stabilisierung des Zustands und Unterstützung im Alltag.')
                ]
            },
            {
                'heading': 'Aufenthaltskonzept Kassel',
                'items': [
                    ('paragraph', 'An der Sommerbergstraße 14 in Kassel bieten wir Patienten eigene Zimmer mit 24-Stunden-Betreuung.'),
                    ('bullet', 'Eigenes Zimmer mit eigenen Möbeln'),
                    ('bullet', 'Alle notwendigen medizinischen Geräte'),
                    ('bullet', 'Großzügiger Garten'),
                    ('bullet', 'Bibliothek und Gemeinschaftsräume'),
                    ('bullet', '5,5 Mitarbeiter pro Patient'),
                    ('bullet', '24h Betreuung und Überwachung'),
                    ('bullet', 'Umfassende pflegerische und medizinische Leistungen'),
                    ('bullet', 'Individuell gestalteter Tagesablauf')
                ]
            },
            {
                'heading': '8 Fachbereiche der Schulung',
                'items': [
                    ('paragraph', 'Regelmäßig durchgeführt unter Leitung von Fachärzten für Pneumologie, Kardiologie und Anästhesiologie.'),
                    ('card', 'Hygiene und Infektionsschutz', 'Desinfektion, Sterilisation, Schutzmaßnahmen und Impfungen.'),
                    ('card', 'Atmung und Beatmung', 'Respiratorische Störungen, Beatmungstechniken, Intubation, Medikamente.'),
                    ('card', 'Herz-Kreislauf-System', 'Herzinsuffizienz, Reanimation, Defibrillation, Medikation.'),
                    ('card', 'Schockgeschehen', '6 Schockformen: hypovolämisch, kardiogen, septisch-toxisch, anaphylaktisch.'),
                    ('card', 'Internistische Notfälle', 'Tracheotomie, Pleurapunktion, PEEP-Beatmung, Druckinfusion.'),
                    ('card', 'Traumatologische Notfälle', 'Schädel-Hirn-Verletzungen, Blutungen, Erstversorgung.'),
                    ('card', 'Brustkorbverletzungen', 'Lungenkontusion, Hämatothorax, Pneumothorax, Herzkontusion.'),
                    ('card', 'Thermische Notfälle', 'Hitze- und Kälteschäden, Verbrennungen, Inhalationstrauma.')
                ]
            },
            {
                'heading': 'Wer die Kosten trägt',
                'items': [
                    ('paragraph', 'Erfahrene Mitarbeiter vom Intensivpflegedienst KSK Farmos GmbH & Co. KG übernehmen nach Erhalt einer Vollmacht die Verhandlungen mit den Kostenträgern bezüglich der Finanzierung einer 24-stündigen Intensiv-Versorgung. Diese sind in der Regel:'),
                    ('card', 'Krankenkasse (SGB V)', 'Anspruch auf häusliche Krankenpflege (Intensivpflege) besteht, wenn eine im Haus lebende Person den Kranken nicht in dem erforderlichen Umfang pflegen und versorgen kann. Die Krankenkasse kann zusätzlich zu den Kosten der Behandlungspflege auch die Kosten der häuslichen Krankenpflege, der Grundpflege und der hauswirtschaftlichen Versorgung übernehmen.'),
                    ('card', 'Pflegekasse (SGB XI)', 'Die Pflegekassen sind für die Sicherstellung der pflegerischen Versorgung ihrer Versicherten verantwortlich. Sie arbeiten mit allen an der pflegerischen und gesundheitlichen Versorgung Beteiligten eng zusammen. Die Pflegekassen stellen insbesondere sicher, dass ärztliche Behandlung, Grundpflege sowie Behandlungspflege und hauswirtschaftliche Versorgung nahtlos und störungsfrei ineinandergreifen.'),
                    ('card', 'Sozialamt', 'Manchmal übersteigen die Kosten des tatsächlichen Pflegebedarfs die Leistungen der oben genannten Kostenträger. In diesem Fall muss der verbleibende Kostenanteil entweder privat getragen werden oder — sofern die Voraussetzungen gegeben sind — von unterschiedlichen Ämtern (Sozialamt, Beihilfe, Regierungspräsidium ...) übernommen werden.'),
                    ('card', 'Private Versicherung', 'Eine zusätzlich abgeschlossene Kranken- oder Unfallversicherung kann je nach Prüfung einen Anteil der anfallenden Kosten übernehmen.')
                ]
            }
        ]
    },
    {
        'title': 'Über uns',
        'desc': 'Unternehmensphilosophie & Team',
        'blocks': [
            {
                'heading': 'Erfahrung und Fürsorge seit 2013',
                'items': [
                    ('paragraph', 'Seit 2013 versorgen wir Menschen mit schweren Erkrankungen — professionell, menschlich und zuverlässig.')
                ]
            },
            {
                'heading': 'Unsere Geschichte — Gegründet aus Überzeugung',
                'items': [
                    ('paragraph', 'Der ambulante Intensivpflegedienst KSK Farmos GmbH & Co. KG wurde 2013 von Viktor Beresnev in Volkmarsen gegründet — mit einem klaren Ziel: Menschen mit schweren Erkrankungen ein würdevolles, selbstbestimmtes Leben zu ermöglichen.'),
                    ('paragraph', 'Seitdem versorgen wir Patienten hessenweit — in ihrem eigenen Zuhause oder in unserer Einrichtung in Kassel.')
                ]
            },
            {
                'heading': 'Unser Leitbild — Woran wir uns orientieren',
                'items': [
                    ('card', 'Lebensprozesse unterstützen', 'Selbstständigkeit fördern und natürliche Prozesse begleiten.'),
                    ('card', 'Körperfunktionen fördern', 'Gezielte Pflege zur Erhaltung körperlicher Funktionen.'),
                    ('card', 'Krankheit kontrollieren', 'Verläufe überwachen, frühzeitig reagieren.'),
                    ('card', 'Wohlbefinden fördern', 'Emotionales und körperliches Wohlbefinden im Fokus.'),
                    ('card', 'Komplikationen verhindern', 'Prophylaxen und kontinuierliche Überwachung.')
                ]
            },
            {
                'heading': 'Die Menschen hinter KSK Farmos GmbH & Co. KG',
                'items': [
                    ('card', 'Pflegedienstleitung (PDL)', 'Verantwortlich für die fachliche Leitung und Qualitätssicherung aller Pflegemaßnahmen.'),
                    ('card', 'Stellv. Pflegedienstleitung', 'Koordination der Pflegeteams und direkte Ansprechpartnerin für Angehörige.'),
                    ('card', 'Unser Pflegeteam', 'Rund 30 Mitarbeiter — qualifiziert, empathisch und zuverlässig im Einsatz.')
                ]
            },
            {
                'heading': 'Unsere Partner & Kooperationen',
                'items': [
                    ('bullet', 'Haus- und Fachärzte in ganz Hessen'),
                    ('bullet', 'Physiotherapeuten'),
                    ('bullet', 'Logopäden'),
                    ('bullet', 'Ergotherapeuten'),
                    ('bullet', 'Medizintechniker'),
                    ('bullet', 'Kranken-, Pflege- und Sozialkassen')
                ]
            },
            {
                'heading': 'Unsere Standorte',
                'items': [
                    ('bullet', 'Hauptzentrale Volkmarsen'),
                    ('bullet', 'Aufenthaltskonzept Kassel')
                ]
            }
        ]
    },
    {
        'title': 'Karriere',
        'desc': 'Stellenangebote & Arbeitgeberleistungen',
        'blocks': [
            {
                'heading': 'Arbeiten wo es zählt.',
                'items': [
                    ('paragraph', 'Übertarifliche Bezahlung, kontinuierliche Weiterbildung und ein Team, das zusammenhält.')
                ]
            },
            {
                'heading': 'Offene Stellen',
                'items': [
                    ('bullet', 'Examinierte Pflegefachkraft (w/m/d)'),
                    ('bullet', 'Pflegekraft / Altenpfleger (w/m/d)'),
                    ('bullet', 'Medizinstudenten — Pflegepraktikum')
                ]
            },
            {
                'heading': 'Was wir Ihnen bieten',
                'items': [
                    ('card', 'Übertarifliche Bezahlung', 'Wir zahlen deutlich über Tarif — weil uns Ihre Arbeit viel wert ist.'),
                    ('card', 'Starkes Team', 'Kollegiale Atmosphäre, gegenseitige Unterstützung und echte Wertschätzung.'),
                    ('card', 'Kontinuierliche Weiterbildung', 'Regelmäßige Fortbildungen durch Fachärzte — auf Kosten des Unternehmens.'),
                    ('card', 'Flexibel und regional', 'Einsätze in ganz Hessen — mit Firmenwagen und flexiblen Arbeitszeiten.'),
                    ('card', 'Sinnvolle Arbeit', 'Sie machen einen echten Unterschied im Leben schwer kranker Menschen.'),
                    ('card', 'Entwicklungsmöglichkeiten', 'Aufstiegschancen und persönliche Weiterentwicklung im Unternehmen.')
                ]
            },
            {
                'heading': 'So bewerben Sie sich',
                'items': [
                    ('card', '1. Bewerbung einreichen', 'Per Formular, E-Mail oder Telefon.'),
                    ('card', '2. Persönliches Gespräch', 'Wir lernen Sie kennen — locker und auf Augenhöhe.'),
                    ('card', '3. Willkommen im Team', 'Einarbeitung, Fortbildungen und Unterstützung.')
                ]
            }
        ]
    },
    {
        'title': 'Schnellbewerbung',
        'desc': 'Interaktiver Bewerbungsprozess',
        'blocks': [
            {
                'heading': 'In 3 Schritten bewerben',
                'items': [
                    ('paragraph', 'Kein Anschreiben nötig — einfach Kontaktdaten hinterlassen und wir melden uns.'),
                    ('card', 'Schritt 1: Über Sie', 'Vorname, Nachname, E-Mail, Telefon'),
                    ('card', 'Schritt 2: Qualifikation', 'Gewünschte Stelle, Arbeitszeit (Vollzeit, Teilzeit, Minijob, Flexibel)'),
                    ('card', 'Schritt 3: Fast geschafft', 'Nachricht / Anschreiben (optional), Datenschutzerklärung zustimmen')
                ]
            }
        ]
    },
    {
        'title': 'FAQ',
        'desc': 'Häufig gestellte Fragen',
        'blocks': [
            {
                'heading': 'Häufig gestellte Fragen rund um unsere Intensivpflege',
                'items': [
                    ('card', 'Ist eine ambulante Intensivpflege zu Hause wirklich möglich?', 'Ja — in vielen Fällen ist eine vollständige Intensivversorgung in den eigenen vier Wänden möglich. Voraussetzung ist eine ärztliche Verordnung häuslicher Krankenpflege.'),
                    ('card', 'In welchen Regionen sind Sie tätig?', 'Wir versorgen Patienten in ganz Hessen — von Volkmarsen und Kassel bis Frankfurt und darüber hinaus.'),
                    ('card', 'Wer übernimmt die Kosten für die Intensivpflege?', 'Die Kosten werden in der Regel von der Kranken- und Pflegekasse übernommen. Wir übernehmen die gesamte Kommunikation mit den Kostenträgern.'),
                    ('card', 'Spielt die Pflegestufe eine Rolle für die 24h-Versorgung?', 'Die 24h-Intensivversorgung wird primär über §37 SGB V finanziert und ist nicht direkt an den Pflegegrad gebunden. Ein höherer Pflegegrad ermöglicht zusätzliche Leistungen.'),
                    ('card', 'Wie lange dauert es, bis die Versorgung beginnt?', 'In der Regel 4 bis 6 Wochen vom Erstgespräch bis zum Versorgungsbeginn.'),
                    ('card', 'Welche Qualifikationen hat Ihr Pflegepersonal?', 'Alle ärztlich verordneten Leistungen werden von examinierten Pflegefachkräften erbracht, die regelmäßig durch Fachärzte fortgebildet werden.'),
                    ('card', 'Was ist das Aufenthaltskonzept in Kassel?', 'An der Sommerbergstraße 14 bieten wir eigene Zimmer, Garten, Bibliothek und 24h-Betreuung mit 5,5 Mitarbeitern pro Patient.'),
                    ('card', 'Wie funktioniert der Wechsel aus dem Krankenhaus nach Hause?', 'Wir kümmern uns um das komplette Überleitungsmanagement. Unser Team koordiniert den reibungslosen Übergang in enger Absprache mit dem Klinikpersonal und organisiert alle benötigten Hilfsmittel.'),
                    ('card', 'Müssen wir medizinische Geräte selbst kaufen?', 'Nein. Alle medizinischen Geräte wie Beatmungsgeräte oder Monitore werden von der Krankenkasse gestellt. Wir arbeiten hier eng mit spezialisierten Medizintechnik-Unternehmen zusammen.'),
                    ('card', 'Sind auch nachts Pflegekräfte vor Ort?', 'Ja. Bei der 24-Stunden-Intensivpflege ist immer eine examinierte Pflegefachkraft bei Ihnen vor Ort — auch nachts —, um eine lückenlose und sichere Überwachung zu gewährleisten.')
                ]
            }
        ]
    },
    {
        'title': 'Kontakt',
        'desc': 'Standorte & Kontaktaufnahme',
        'blocks': [
            {
                'heading': 'Sprechen Sie uns an — Kostenlos. Unverbindlich. Menschlich.',
                'items': [
                    ('bullet', 'Zentrale Volkmarsen: 05693 / 9189907'),
                    ('bullet', 'Mobil / 24h-Bereitschaft: 0170 / 7652593'),
                    ('bullet', 'Faxnummer: 05693 / 9189908'),
                    ('bullet', 'E-Mail: pflege@ksk-farmos.de')
                ]
            }
        ]
    },
    {
        'title': 'Beratung',
        'desc': 'Kostenloses Erstgespräch anfordern',
        'blocks': [
            {
                'heading': 'Kostenlose Erstberatung',
                'items': [
                    ('paragraph', 'Wir nehmen uns Zeit für Sie — persönlich, telefonisch oder per E-Mail.'),
                    ('paragraph', 'Rückmeldung innerhalb 24 Stunden — garantiert. Wir klären alle Ihre Fragen zur 24h-Intensivpflege und der Finanzierung.')
                ]
            }
        ]
    }
]

# ─────────────────────────────────────────────────────────
# DATA 2: KSK PATIENTS LANDING PAGE
# ─────────────────────────────────────────────────────────
pat_pages = [
    {
        'title': 'Patienten-Landingpage',
        'desc': 'Ambulante Intensivpflege Nordhessen',
        'blocks': [
            {
                'heading': 'Ihr Angehöriger braucht Pflege? Wir sind für Sie da.',
                'items': [
                    ('paragraph', 'Kostenlose Erstberatung zur ambulanten Intensivpflege in Nordhessen. Wir klären die Finanzierung — die Krankenkasse übernimmt die Kosten.'),
                    ('paragraph', 'Die Krankenkasse übernimmt die Kosten für die Intensivpflege in der Regel vollständig. Wir kümmern uns um die gesamte Kommunikation mit den Kostenträgern — Sie müssen sich um nichts kümmern.')
                ]
            },
            {
                'heading': 'Was wir für Sie tun (Wir nehmen Ihnen alles ab)',
                'items': [
                    ('bullet', 'Kostenlose Erstberatung — unverbindlich und persönlich'),
                    ('bullet', 'Kostenübernahme durch Krankenkasse — wir übernehmen die Kommunikation'),
                    ('bullet', 'Pflege bei Ihnen zu Hause — oder in unserer Einrichtung in Kassel'),
                    ('bullet', 'Examinierte Pflegefachkräfte — 5,5 Mitarbeiter pro Patient'),
                    ('bullet', 'Rückmeldung innerhalb 24 Stunden — garantiert'),
                    ('bullet', 'Versorgung in ganz Hessen — seit 2013'),
                    ('quote', '"Das Wohl des Patienten ist oberstes Gesetz." — Unser Leitbild')
                ]
            },
            {
                'heading': 'Welche Patienten wir versorgen (Diagnosen)',
                'items': [
                    ('card', 'Beatmungspatienten', 'Invasive und nicht-invasive Beatmung'),
                    ('card', 'Wachkoma (Apallisches Syndrom)', 'Langzeitversorgung und Stimulation'),
                    ('card', 'ALS', 'Progressive Begleitung bei ALS und neurologischen Erkrankungen'),
                    ('card', 'Neurologische Erkrankungen', 'Ganzheitliche Pflege und Unterstützung'),
                    ('card', 'Querschnittslähmung', 'Spezialisierte Pflege bei Rückenmarksverletzungen'),
                    ('card', 'Schädel-Hirn-Trauma', 'Rehabilitation, Pflege und langfristige Stabilisierung')
                ]
            },
            {
                'heading': 'In 3 Schritten zur Versorgung',
                'items': [
                    ('card', '1. Kostenloses Erstgespräch', 'Wir lernen Ihre Situation kennen — telefonisch oder persönlich.'),
                    ('card', '2. Individuelle Planung', 'Wir erstellen den Pflegeplan und klären die Kostenübernahme mit der Kasse.'),
                    ('card', '3. Versorgungsbeginn', 'Ihr persönliches Pflegeteam beginnt die Betreuung — zuverlässig und menschlich.')
                ]
            },
            {
                'heading': 'Jetzt kostenlos anfragen',
                'items': [
                    ('paragraph', 'Wir sind für Sie da — persönlich, schnell und ohne Verpflichtung.')
                ]
            },
            {
                'heading': 'Erfahren Sie mehr über uns',
                'items': [
                    ('paragraph', 'Besuchen Sie unsere Hauptseite für ausführliche Informationen zu unseren Leistungen, dem Team und unserer Philosophie.')
                ]
            }
        ]
    }
]

if __name__ == '__main__':
    # Build Corporate
    build_docx(
        filepath="/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-farmos/site/сайт-текст/KSK-Farmos-Website-Text-DE.docx",
        doc_title="KSK Farmos GmbH & Co. KG",
        doc_subtitle="Inhalt der Website (Deutsch / Original)",
        metadata_text="Dokumententyp: Website-Textgliederung (ohne Menü und Fußzeile)\nGeneriert am: 18. Mai 2026\nIntensivpflegedienst Nordhessen",
        pages=corp_pages
    )

    # Build Patients
    build_docx(
        filepath="/Users/skaissts/Desktop/macbook/работа/antigravity/ksk-farmos/site/сайт-текст/KSK-Patients-Landing-Text-DE.docx",
        doc_title="KSK Farmos — Intensivpflege",
        doc_subtitle="Inhalt der Patienten-Landingpage (Deutsch / Original)",
        metadata_text="Dokumententyp: Landingpage-Textgliederung (ohne Menü und Fußzeile)\nGeneriert am: 18. Mai 2026\nPatienten-Werbelandingpage Nordhessen",
        pages=pat_pages
    )
