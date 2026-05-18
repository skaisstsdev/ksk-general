import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    # Paths
    base_dir = os.path.dirname(os.path.abspath(__file__))
    de_json_path = os.path.join(base_dir, '../i18n/de.json')
    ru_json_path = os.path.join(base_dir, '../i18n/ru.json')
    md_path = os.path.join(base_dir, 'volltext-de.md')
    output_docx_path = os.path.join(base_dir, 'KSK-Farmos-Website-Text.docx')

    # Load translations
    with open(de_json_path, 'r', encoding='utf-8') as f:
        de_dict = json.load(f)

    with open(ru_json_path, 'r', encoding='utf-8') as f:
        ru_dict = json.load(f)

    # Read the markdown source structure
    with open(md_path, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    doc = Document()

    # Style definitions
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color palette
    c_violet = RGBColor(76, 43, 115)      # #4C2B73 - Quiet Luxury Deep Violet
    c_charcoal = RGBColor(44, 44, 44)     # #2C2C2C - Dark Charcoal
    c_muted = RGBColor(110, 110, 110)     # #6E6E6E - Muted Gray

    # ─────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────
    # Spacing at top
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(80)

    # Brand Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("KSK Farmos GmbH & Co. KG")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = c_violet
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(10)

    # Brand Subtitle
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Полное текстовое содержание сайта и переводы")
    run_sub.font.name = 'Georgia'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = c_charcoal
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)

    p_sub2 = doc.add_paragraph()
    run_sub2 = p_sub2.add_run("Volles Textinhalt der Website und Übersetzungen")
    run_sub2.font.name = 'Georgia'
    run_sub2.font.size = Pt(12)
    run_sub2.font.italic = True
    run_sub2.font.color.rgb = c_muted
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_after = Pt(180)

    # Metadata
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("Составитель: Antigravity AI Assistant\nДата создания: 18 мая 2026 г.\nЯзыковые версии: Русский / Deutsch\nСлужба интенсивного ухода (Kassel / Volkmarsen)")
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = c_muted
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # End of cover page
    doc.add_page_break()

    # Helper lookup function
    def get_ru_translation(de_line):
        clean_de = de_line.strip().strip('-').strip().strip('"').strip("'").strip()
        if not clean_de:
            return ""
        
        # Exact lookup
        for k, v in de_dict.items():
            if v.strip() == clean_de:
                return ru_dict.get(k, clean_de)
            if v.strip().lower() == clean_de.lower():
                return ru_dict.get(k, clean_de)
                
        # Structural fallback translations
        fallback_map = {
            "Header": "Шапка сайта (Header)",
            "Footer": "Подвал сайта (Footer)",
            "Language Modal": "Окно выбора языка (Language Switcher)",
            "Back to Top": "Кнопка «Наверх»",
            "Hero": "Главный баннер (Hero)",
            "Editorial Accent": "Цитата / Вводное слово",
            "Service Cards": "Карточки услуг (Направления помощи)",
            "Diagnosen": "Диагнозы (Кого мы принимаем на попечение)",
            "3 Steps": "3 простых шага к началу ухода",
            "Über uns Preview": "Краткое превью «О нас»",
            "Dual CTA": "Двойной призыв к действию (CTA для пациентов и соискателей)",
            "Anchor Nav": "Якорное меню переходов",
            "Häusliche Pflege": "Раздел: Интенсивный уход на дому",
            "Beatmung Highlight": "Ключевая компетенция: Интенсивный уход на ИВЛ",
            "Aufenthaltskonzept": "Раздел: Концепция проживания в Касселе",
            "Fortbildungen": "Раздел: Обучение и 8 направлений повышения квалификации",
            "Kostenlose Beratung": "Форма заказа бесплатной консультации",
            "Kostenübernahme": "Финансирование и покрытие расходов больничными кассами",
            "CTA": "Блок призыва к действию",
            "Geschichte": "Наша история и дата основания",
            "Leitbild": "Наши ориентиры и миссия ухода",
            "Team": "Руководство и команда по уходу",
            "Partner": "Наши партнеры и кооперации",
            "Standorte": "Наши адреса и филиалы",
            "Vorteile": "Преимущества работы у нас",
            "Stellenangebote": "Открытые вакансии",
            "Bewerbung Steps": "Простые этапы трудоустройства",
            "Wizard Form": "Интерактивная анкета быстрого отклика (Шаги)",
            "Search & Categories": "Поиск и фильтр категорий вопросов",
            "Accordion": "Вопросы и подробные ответы (Аккордеон)",
            "Top Cards": "Важные факты (Стоимость, Сроки, Квалификация)",
            "Office Cards": "Карточки филиалов и офисов",
            "Form": "Форма обратной связи",
            "Bewerber": "Блок для соискателей",
            "Split": "Информационный сплит-блок",
            "ОБЩИЕ ЭЛЕМЕНТЫ (Header / Footer / Lang / CTA)": "ОБЩИЕ СИСТЕМНЫЕ ЭЛЕМЕНТЫ (Шапка / Подвал / Язык)",
            "INDEX.HTML — Startseite": "ГЛАВНАЯ СТРАНИЦА (index.html)",
            "LEISTUNGEN.HTML": "СТРАНИЦА УСЛУГ (leistungen.html)",
            "UEBER-UNS.HTML": "СТРАНИЦА «О НАС» (ueber-uns.html)",
            "KARRIERE.HTML": "СТРАНИЦА КАРЬЕРЫ И ВАКАНСИЙ (karriere.html)",
            "FAQ.HTML": "СТРАНИЦА ВОПРОСОВ И ОТВЕТОВ (faq.html)",
            "KONTAKT.HTML": "СТРАНИЦА КОНТАКТОВ (kontakt.html)",
            "BERATUNG.HTML": "СТРАНИЦА ЗАПИСИ НА КОНСУЛЬТАЦИЮ (beratung.html)",
            "IMPRESSUM.HTML": "ВЫХОДНЫЕ ДАННЫЕ (impressum.html)",
            "DATENSCHUTZ.HTML": "ПОЛИТИКА КОНФИДЕНЦИАЛЬНОСТИ (datenschutz.html)"
        }
        for k, v in fallback_map.items():
            if k.lower() in de_line.lower():
                return v
        return clean_de

    # ─────────────────────────────────────────────────────────
    # PART 1: RUSSIAN TRANSLATION (РУССКАЯ ВЕРСИЯ)
    # ─────────────────────────────────────────────────────────
    p_part1 = doc.add_paragraph()
    run_part1 = p_part1.add_run("ЧАСТЬ 1. РУССКАЯ ВЕРСИЯ САЙТА (ПЕРЕВОД)")
    run_part1.font.name = 'Georgia'
    run_part1.font.size = Pt(20)
    run_part1.font.bold = True
    run_part1.font.color.rgb = c_violet
    p_part1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_part1.paragraph_format.space_before = Pt(20)
    p_part1.paragraph_format.space_after = Pt(20)

    for line in md_lines:
        line_str = line.strip()
        if not line_str:
            continue
        
        # Main Page Heading (H2 in markdown -> Heading 1 in Word)
        if line_str.startswith("## ") and not line_str.startswith("###"):
            title_text = line_str.replace("## ", "")
            ru_title = get_ru_translation(title_text)
            
            doc.add_page_break()
            
            p = doc.add_paragraph()
            run = p.add_run(ru_title.upper())
            run.font.name = 'Georgia'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = c_violet
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            
        # Section Heading (H3 in markdown -> Heading 2 in Word)
        elif line_str.startswith("### "):
            section_text = line_str.replace("### ", "")
            ru_section = get_ru_translation(section_text)
            
            p = doc.add_paragraph()
            run = p.add_run(f"Блок: {ru_section}")
            run.font.name = 'Georgia'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = c_charcoal
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        # Text Line (List items in markdown -> Paragraphs in Word)
        elif line_str.startswith("- "):
            text_val = line_str.replace("- ", "")
            ru_text = get_ru_translation(text_val)
            
            # Formatting quotes specially
            if ru_text.startswith('"') or ru_text.startswith('«'):
                p = doc.add_paragraph()
                run = p.add_run(f"   {ru_text}")
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.italic = True
                run.font.color.rgb = c_charcoal
                p.paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"•  {ru_text}")
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.color.rgb = c_charcoal
                p.paragraph_format.space_after = Pt(4)
                
        elif line_str == "---":
            # Add subtle divider
            p = doc.add_paragraph()
            run = p.add_run("____________________________________________________")
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = c_muted
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)

    # ─────────────────────────────────────────────────────────
    # PART 2: GERMAN ORIGINAL (DEUTSCHE ORIGINALVERSION)
    # ─────────────────────────────────────────────────────────
    doc.add_page_break()
    
    p_part2 = doc.add_paragraph()
    run_part2 = p_part2.add_run("TEIL 2. DEUTSCHE ORIGINALVERSION DER WEBSITE")
    run_part2.font.name = 'Georgia'
    run_part2.font.size = Pt(20)
    run_part2.font.bold = True
    run_part2.font.color.rgb = c_violet
    p_part2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_part2.paragraph_format.space_before = Pt(20)
    p_part2.paragraph_format.space_after = Pt(20)

    for line in md_lines:
        line_str = line.strip()
        if not line_str:
            continue
        
        # Main Page Heading
        if line_str.startswith("## ") and not line_str.startswith("###"):
            title_text = line_str.replace("## ", "")
            
            doc.add_page_break()
            
            p = doc.add_paragraph()
            run = p.add_run(title_text.upper())
            run.font.name = 'Georgia'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = c_violet
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            
        # Section Heading
        elif line_str.startswith("### "):
            section_text = line_str.replace("### ", "")
            
            p = doc.add_paragraph()
            run = p.add_run(f"Bereich: {section_text}")
            run.font.name = 'Georgia'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = c_charcoal
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        # Text Line
        elif line_str.startswith("- "):
            text_val = line_str.replace("- ", "")
            
            if text_val.startswith('"'):
                p = doc.add_paragraph()
                run = p.add_run(f"   {text_val}")
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.italic = True
                run.font.color.rgb = c_charcoal
                p.paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"•  {text_val}")
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.color.rgb = c_charcoal
                p.paragraph_format.space_after = Pt(4)
                
        elif line_str == "---":
            p = doc.add_paragraph()
            run = p.add_run("____________________________________________________")
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = c_muted
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)

    # Save
    doc.save(output_docx_path)
    print(f"Document successfully created at: {output_docx_path}")

if __name__ == '__main__':
    create_document()
